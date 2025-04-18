from fastapi import FastAPI, UploadFile, File
from fastapi.responses import HTMLResponse, JSONResponse, FileResponse, StreamingResponse
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
import shutil
import cv2
import easyocr
import torch
import re
import fitz 
from PIL import Image
from vietocr.tool.predictor import Predictor
from vietocr.tool.config import Cfg
from docx import Document
from io import BytesIO
import os
import json
from fastapi import Request
from pydantic import BaseModel
from fastapi.responses import StreamingResponse
from concurrent.futures import ThreadPoolExecutor, as_completed
from tempfile import NamedTemporaryFile
import zipfile
import asyncio
import time

from extract_data import extract_data
from draw_bounding_boxes import draw_bounding_boxes, process_outside_table_images
from save_json_words import save_json, save_word
from full_table import full_pipeline


app = FastAPI()

# Cấu hình template HTML
templates = Jinja2Templates(directory="templates")

# Cấu hình thư mục tĩnh để lưu ảnh tải lên
app.mount("/static", StaticFiles(directory="static"), name="static")

# Tạo thư mục lưu trữ nếu chưa có
os.makedirs("static/uploaded_images", exist_ok=True)
os.makedirs("static/processed_images", exist_ok=True)
os.makedirs("static/output_json", exist_ok=True)
os.makedirs("static/output_docx", exist_ok=True)
os.makedirs("static/output_png", exist_ok=True)

# Route stream trả kết quả theo từng dòng cho (Text only)
@app.post("/stream_process/")
async def stream_process(file: UploadFile = File(...)):
    # Lưu file vào temp file trước khi xử lý
    with NamedTemporaryFile(delete=False) as temp_file:
        shutil.copyfileobj(file.file, temp_file)
        temp_path = temp_file.name

    async def generate():
        try:
            # Gọi hàm xử lý ảnh của bạn
            extracted_texts, _ = draw_bounding_boxes(temp_path)
            
            # Stream từng dòng kết quả
            for idx, text in enumerate(extracted_texts):
                yield json.dumps({
                    "text": text,
                }) + "\n"
                
                await asyncio.sleep(0.1)  # Điều chỉnh tốc độ hiển thị
                
        finally:
            # Dọn dẹp temp file
            if os.path.exists(temp_path):
                os.unlink(temp_path)

    return StreamingResponse(generate(), media_type="text/event-stream")

# Route stream trả kết quả theo từng dòng (Table included)
@app.post("/stream_process_table/")
async def stream_process_table(file: UploadFile = File(...)):
    # ======== BƯỚC 0: Làm sạch các thư mục tạm =========
    folders_to_clean = [
        "CROP_TABLE/line_paddle",
        "CROP_TABLE/output_txt",
        "CROP_TABLE/output_cells",
        "CROP_TABLE/outside_table",
        "CROP_TABLE/result_json",
        "CROP_TABLE/tables"
    ]

    for folder in folders_to_clean:
        if os.path.exists(folder):
            for f in os.listdir(folder):
                f_path = os.path.join(folder, f)
                if os.path.isfile(f_path):
                    os.remove(f_path)
                elif os.path.isdir(f_path):
                    shutil.rmtree(f_path)
            print(f"🧹 Đã xoá thư mục tạm: {folder}")

    # ======== BƯỚC 1: Lưu ảnh upload vào thư mục tạm =========
    temp_upload_dir = "temp_upload"
    os.makedirs(temp_upload_dir, exist_ok=True)
    image_path = os.path.join(temp_upload_dir, file.filename)

    with open(image_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    # ======== BƯỚC 2: Gọi pipeline xử lý bảng =========
    output_dir = "CROP_TABLE"
    vietocr_weight_path = "transformerocr(6).pth"
    template_path = os.path.join(output_dir, "table_sample.json")

    json_output_path = full_pipeline(
        image_path=image_path,
        output_dir=output_dir,
        vietocr_weight_path=vietocr_weight_path,
        template_path=template_path
    )

    # ======== BƯỚC 3: Stream kết quả đúng thứ tự vùng ngoài bảng & bảng =========
    async def generate():
        order_path = os.path.join(output_dir, "order_list.json")
        if not os.path.exists(order_path):
            yield json.dumps({"type": "error", "text": "Không tìm thấy thứ tự ảnh!"}) + "\n"
            return

        with open(order_path, "r", encoding="utf-8") as f:
            order_list = json.load(f)

        for item in order_list:
            if item["type"] == "text":
                image_path = item["path"]
                if os.path.exists(image_path):
                    lines, _ = draw_bounding_boxes(image_path)
                    for text in lines:
                        yield json.dumps({"type": "text", "text": text}) + "\n"
                        await asyncio.sleep(0.1)
            elif item["type"] == "table":
                table_id = item["id"]
                table_json_path = os.path.join(output_dir, "result_json", f"{table_id}_output.json")
                if os.path.exists(table_json_path):
                    with open(table_json_path, "r", encoding="utf-8") as f:
                        table_data = json.load(f)
                        yield json.dumps({
                            "type": "table",
                            "table_id": table_id,
                            "data": table_data
                        }) + "\n"
                        await asyncio.sleep(0.2)

    return StreamingResponse(generate(), media_type="text/event-stream")


# Định nghĩa class OCRText để nhận JSON từ request
class OCRText(BaseModel):
    text: str

# Route xuất dữ liệu JSON
@app.post("/export_json/")
async def export_json(data: dict):
    image_name = data.get("image_name", "output").split(".")[0]
    output_dir = "static/output_json"
    os.makedirs(output_dir, exist_ok=True)
    file_path = os.path.join(output_dir, f"{image_name}.json")

    # === Trường hợp Text Only: gộp tất cả key-value vào 1 object duy nhất ===
    if "text" in data:
        text_content = data["text"]
        json_data = extract_data(text_content)

        with open(file_path, "w", encoding="utf-8") as f:
            json.dump(json_data, f, indent=2, ensure_ascii=False)

        return {
            "file_path": file_path,
            "file_name": f"{image_name}.json"
        }

    # === Trường hợp Table Included ===
    items = data.get("data", [])
    combined_result = {}

    for item in items:
        if item.get("type") == "text":
            text_line = item.get("text", "")
            parsed = extract_data(text_line)
            for k, v in parsed.items():
                if k not in combined_result:
                    combined_result[k] = v

        elif item.get("type") == "table":
            table_data = item.get("data", [])

            # ✅ Replace "" with null
            for row in table_data:
                for k, v in row.items():
                    if v == "":
                        row[k] = None

            combined_result.setdefault("tables", []).append(table_data)

    with open(file_path, "w", encoding="utf-8") as f:
        json.dump(combined_result, f, indent=2, ensure_ascii=False)

    return {
        "file_path": file_path,
        "file_name": f"{image_name}.json"
    }

# Route xuất dữ liệu Word
@app.post("/export_docx/")
async def export_docx(data: dict):
    image_name = data.get("image_name", "output").split(".")[0]
    items = data.get("data", [])

    file_path = os.path.join("static/output_docx", f"{image_name}.docx")
    os.makedirs(os.path.dirname(file_path), exist_ok=True)

    doc = Document()

    for item in items:
        if item.get("type") == "text":
            text = item.get("text", "").strip()
            if text:
                doc.add_paragraph(text)

        elif item.get("type") == "table":
            table_data = item.get("data", [])
            if not table_data:
                continue

            headers = list(table_data[0].keys())
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"

            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header

            for row in table_data:
                row_cells = table.add_row().cells
                for i, key in enumerate(headers):
                    row_cells[i].text = str(row.get(key, ""))

            doc.add_paragraph("")  # Dòng trống sau bảng

    doc.save(file_path)

    return {
        "file_path": file_path,
        "file_name": f"{image_name}.docx"
    }


# === Route OCR Processing ===
@app.post("/process-table-image/")
async def process_table_image(file: UploadFile = File(...)):
    try:
        # ======= 1. Làm sạch thư mục tạm (temp_upload và temp) =======
        for temp_folder in ["temp_upload", "temp"]:
            os.makedirs(temp_folder, exist_ok=True)
            for f in os.listdir(temp_folder):
                f_path = os.path.join(temp_folder, f)
                if os.path.isfile(f_path):
                    os.remove(f_path)

        # ======= 2. Lưu file ảnh mới =======
        image_path = os.path.join("temp_upload", file.filename)
        with open(image_path, "wb") as f:
            shutil.copyfileobj(file.file, f)

        # ======= 3. Gọi pipeline xử lý bảng =========
        output_dir = "CROP_TABLE"
        vietocr_weight_path = "transformerocr(6).pth"
        template_path = os.path.join(output_dir, "table_sample.json")

        json_output_path = full_pipeline(
            image_path=image_path,
            output_dir=output_dir,
            vietocr_weight_path=vietocr_weight_path,
            template_path=template_path
        )

        # ======= 4. Xử lý vùng ngoài bảng =======
        outside_folder = os.path.join(output_dir, "outside_table")
        outside_text = []
        if os.path.exists(outside_folder):
            outside_text = process_outside_table_images(outside_folder)

        # ======= 5. Trả kết quả =======
        return JSONResponse(content={
            "message": "Xử lý thành công!",
            "json_output_path": json_output_path,
            "outside_text": outside_text
        })

    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)

# TRANG PDF TO PNG (Bắt đầu từ đây)
# === Route cho trang OCR Processing ===
@app.get("/", response_class=HTMLResponse)
async def ocr_home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

# === Route cho trang PDF to PNG (hiển thị form) ===
@app.get("/pdf2png", response_class=HTMLResponse)
async def pdf2png_home(request: Request):
    return templates.TemplateResponse("pdf2png.html", {"request": request})

# === Route xử lý file PDF và trả về file ZIP hoặc PNG ===
@app.post("/convert_pdf")
async def convert(file: UploadFile = File(...)):
    # Thêm timestamp để tránh trùng tên file
    timestamp = str(int(time.time()))
    base_name = os.path.splitext(file.filename)[0]
    temp_pdf_path = f"static/output_png/temp_{timestamp}_{file.filename}"
    
    with open(temp_pdf_path, "wb") as f:
        f.write(await file.read())

    try:
        doc = fitz.open(temp_pdf_path)
        output_dir = "static/output_png"

        if len(doc) == 1:
            # Xử lý ảnh đơn trang
            pix = doc.load_page(0).get_pixmap(dpi=200, alpha=False)
            output_filename = os.path.join(output_dir, f"{base_name}_{timestamp}.png")
            
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            img.save(output_filename, "PNG", dpi=(200, 200), optimize=True)
            
            return FileResponse(
                output_filename,
                media_type="image/png",
                headers={"Content-Disposition": f'attachment; filename="{os.path.basename(output_filename)}"'}
            )
        else:
            # Xử lý nhiều trang
            zip_filename = os.path.join(output_dir, f"{base_name}_{timestamp}.zip")
            with zipfile.ZipFile(zip_filename, "w") as zipf:
                for i in range(len(doc)):
                    pix = doc.load_page(i).get_pixmap(dpi=200, alpha=False)
                    img_filename = f"{base_name}_page_{i+1}.png"
                    
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    with BytesIO() as buffer:
                        img.save(buffer, "PNG", dpi=(200, 200), optimize=True)
                        zipf.writestr(img_filename, buffer.getvalue())
            
            return FileResponse(
                zip_filename,
                media_type="application/zip",
                headers={"Content-Disposition": f'attachment; filename="{os.path.basename(zip_filename)}"'}
            )

    except Exception as e:
        return {"error": str(e)}
    
    finally:
        if 'doc' in locals(): doc.close()
        # Xóa file PDF tạm sau khi xử lý
        if os.path.exists(temp_pdf_path): os.remove(temp_pdf_path)

# uvicorn app:app --reload

# Neu load mai khong len web thi bat cmd admin:
# netstat -ano | findstr :8000
# taskkill /PID 21476 /F 